import os
import sys
import subprocess
import threading
import time
import socket
import webbrowser
import shutil
import json
import re
import copy
import uuid
import tempfile
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox

import customtkinter as ctk
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor, Emu
import win32com.client

import socket

def get_local_ip():
    try:
        with socket.socket(socket.AF_INET, socket.SOCK_DGRAM) as s:
            s.connect(("8.8.8.8", 80))
            return s.getsockname()[0]
    except Exception:
        return "127.0.0.1"

# ===================== CẤU HÌNH =====================
DATA_DIR = 'data'
QUESTIONS_FILE = os.path.join(DATA_DIR, 'questions.xlsx')
os.makedirs(DATA_DIR, exist_ok=True)

# ===================== CẤU HÌNH =====================
QUESTION_START_RE = re.compile(r'^\s*(?:Câu\s*(\d+)[\.:]\s*|(\d+)[\.:]\s+)', re.IGNORECASE)
#MCQ_OPTION_RE = re.compile(r'^\s*([A-D])[\.\)]\s*(.*)')
#TF_OPTION_RE = re.compile(r'^\s*([a-d])[\.\)]\s*(.*)')
MCQ_OPTION_RE = re.compile(r'^\s*([A-D])[\.\)](?:\s+|$)')
TF_OPTION_RE = re.compile(r'^\s*([a-d])[\.\)](?:\s+|$)')
ANSWER_LINE_RE = re.compile(r'(?:Đáp\s*án|ĐA|ANSWER)\s*[:;]\s*([A-Da-d])', re.IGNORECASE)
ANSWER_BLOCK_RE = re.compile(r'^(?:Đáp\s*án|ĐA|ANSWER)\s*[:;]\s*$', re.IGNORECASE)
TF_ANSWER_RE = re.compile(r'^\s*([a-d])[\.\)]\s*(Đúng|Sai|True|False)', re.IGNORECASE)
LEVEL_RE = re.compile(r'(Mức|Mức\s*độ|Level)\s*[:\-]?\s*\d*\s*[\(\（]?\s*(biết|hiểu|vận\s*dụng|vd|thông\s*hiểu|nhận\s*biết)[\)\）]?', re.IGNORECASE)
    
# ===================== XỬ LÝ ẢNH =====================
def generate_unique_id():
    time_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_str = str(uuid.uuid4())[:8]
    return f"{time_str}_{unique_str}"

def save_image_safely(image_bytes, output_folder, content_type):
    ext = content_type.split('/')[-1] if '/' in content_type else 'png'
    if ext == 'jpeg': ext = 'jpg'
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    img_id = generate_unique_id()
    filename = f"img_{img_id}.{ext}"
    with open(os.path.join(output_folder, filename), 'wb') as f:
        f.write(image_bytes)
    return filename

def extract_images_to_placeholder(doc, image_folder="extracted_images"):
    print(f"\n📸 Đang xử lý ảnh trong file...")
    image_count = 0
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    
    for para in doc.paragraphs:
        for run in para.runs:
            drawings = run._element.findall(qn('w:drawing'))
            for drawing in drawings:
                blips = drawing.xpath('.//a:blip')
                extents = drawing.xpath('.//wp:extent')
                for idx, blip in enumerate(blips):
                    rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rid:
                        try:
                            image_part = para.part.related_parts[rid]
                            saved_name = save_image_safely(image_part.blob, image_folder, image_part.content_type)
                            width = height = 0
                            if idx < len(extents):
                                cx = extents[idx].get('cx')
                                cy = extents[idx].get('cy')
                                if cx and cy:
                                    width = int(cx) if cx else 0
                                    height = int(cy) if cy else 0
                            
                            placeholder = f" [IMG: {saved_name} | {width}x{height}] "
                            t_element = OxmlElement('w:t')
                            t_element.set(qn('xml:space'), 'preserve')
                            t_element.text = placeholder
                            drawing.addprevious(t_element)
                            parent = drawing.getparent()
                            if parent is not None: 
                                parent.remove(drawing)
                            image_count += 1
                        except Exception as e: 
                            continue
    print(f"  ✅ Đã xử lý {image_count} ảnh")
    return doc

def restore_images_in_document(doc, image_folder):
    pattern = re.compile(r"\[IMG:\s*(.+?)\s*\|\s*(\d+)x(\d+)\]")
    print(f"  🔍 Đang phục hồi ảnh...")
    image_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if not text: continue
            matches = list(pattern.finditer(text))
            if not matches: continue
            
            run_bold = run.bold
            run_italic = run.italic
            run_underline = run.underline
            run_font_size = run.font.size
            run_font_color = run.font.color.rgb if run.font.color else None
            
            run.text = ""
            last_pos = 0
            for match in matches:
                before_text = text[last_pos:match.start()]
                if before_text:
                    new_run = para.add_run(before_text)
                    new_run.bold = run_bold
                    new_run.italic = run_italic
                    new_run.underline = run_underline
                    if run_font_size: new_run.font.size = run_font_size
                    if run_font_color: new_run.font.color.rgb = run_font_color
                
                filename = match.group(1).strip()
                width = int(match.group(2))
                height = int(match.group(3))
                img_path = os.path.join(image_folder, filename)
                
                if os.path.exists(img_path):
                    try:
                        img_run = para.add_run()
                        if width > 0 and height > 0:
                            img_run.add_picture(img_path, width=Emu(width), height=Emu(height))
                        else:
                            img_run.add_picture(img_path)
                        image_count += 1
                    except:
                        err_run = para.add_run(f"[LỖI: {filename}]")
                        err_run.bold = True
                        err_run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    err_run = para.add_run(f"[MẤT ẢNH: {filename}]")
                    err_run.bold = True
                    err_run.font.color.rgb = RGBColor(255, 0, 0)
                
                last_pos = match.end()
            
            after_text = text[last_pos:]
            if after_text:
                new_run = para.add_run(after_text)
                new_run.bold = run_bold
                new_run.italic = run_italic
                new_run.underline = run_underline
                if run_font_size: new_run.font.size = run_font_size
                if run_font_color: new_run.font.color.rgb = run_font_color
    print(f"  ✅ Đã phục hồi {image_count} ảnh")
    return doc

# ===================== HÀM PHỤ =====================
def convert_numbering_to_text_with_word(input_path, output_path):
    print("\n🔄 Đang chuẩn hóa văn bản và ngắt dòng...")
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(os.path.abspath(input_path))
        
        # Chuyển numbering thành văn bản thuần
        doc.ConvertNumbersToText()
        
        # Thay thế Manual Line Break (^l) thành Paragraph Mark (^p)
        # Sử dụng phương thức Execute chi tiết để tránh lỗi tùy chọn phiên bản Word
        find_obj = doc.Content.Find
        find_obj.ClearFormatting()
        find_obj.Replacement.ClearFormatting()
        find_obj.Execute(
            FindText="^l", 
            ReplaceWith="^p", 
            Replace=2,              # 2 = wdReplaceAll
            Forward=True, 
            Wrap=1,                 # 1 = wdFindContinue
            Format=False, 
            MatchCase=False, 
            MatchWholeWord=False
        )
        
        doc.SaveAs2(os.path.abspath(output_path), 16)
        doc.Close()
        word.Quit()
        return True
    except Exception as e:
        print(f"❌ Lỗi xử lý Word: {e}")
        try: word.Quit()
        except: pass
        return False
    
def run_is_underlined(run):
    if not hasattr(run, 'font'): return False
    return run.font.underline not in [None, False]

def run_has_red_color(run):
    try:
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            if hasattr(rgb, 'rgb'):
                r, g, b = rgb.rgb[0], rgb.rgb[1], rgb.rgb[2]
            else:
                hex_str = str(rgb).lstrip('#')
                r = int(hex_str[0:2], 16)
                g = int(hex_str[2:4], 16)
                b = int(hex_str[4:6], 16)
            if r > 200 and g < 100 and b < 100:
                return True
        if run.font.highlight_color and run.font.highlight_color in [6, 13, 16]:
            return True
    except:
        pass
    return False

def find_underlined_option_in_paragraph(para):
    if not hasattr(para, 'runs') or not para.runs: return None
    text = para.text.strip()
    match = MCQ_OPTION_RE.match(text) or TF_OPTION_RE.match(text)
    if match:
        label = match.group(1).upper()
        if any(run_is_underlined(run) for run in para.runs if run.text.strip()):
            return label
    return None

def find_red_option_in_paragraph(para):
    if not hasattr(para, 'runs') or not para.runs: return None
    text = para.text.strip()
    match = MCQ_OPTION_RE.match(text) or TF_OPTION_RE.match(text)
    if match:
        label = match.group(1).upper()
        if any(run_has_red_color(run) for run in para.runs if run.text.strip()):
            return label
    return None

def is_run_marked_as_answer(run):
    if not run.text.strip(): return False
    if hasattr(run.font, 'underline') and run.font.underline not in [None, False]: return True
    if hasattr(run.font, 'highlight_color') and run.font.highlight_color is not None: return True
    try:
        if run.font.color and run.font.color.rgb:
            rgb = str(run.font.color.rgb).replace('#', '')
            if len(rgb) >= 6:
                r, g, b = int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16)
                if r > 150 and g < 100 and b < 100: return True
    except: pass
    xml_str = run._element.xml
    if '<w:u ' in xml_str or '<w:highlight' in xml_str: return True
    match = re.search(r'<w:color w:val="([0-9A-Fa-f]{6})"', xml_str)
    if match:
        hex_color = match.group(1)
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        if r > 150 and g < 100 and b < 100: return True
    return False

def remove_level_lines(doc):
    print("\n🗑️ Đang xóa dòng mức độ...")
    count = 0
    for para in list(doc.paragraphs):
        if LEVEL_RE.search(para.text):
            try:
                para._element.getparent().remove(para._element)
                count += 1
            except: pass
    print(f"✅ Đã xóa {count} dòng")
    return doc

def remove_answer_table(doc):
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) == 2:
            first_row = table.rows[0].cells
            if first_row[0].text.strip() == "Câu" and first_row[1].text.strip() == "Đáp án":
                table._element.getparent().remove(table._element)
                break
    return doc

# ===================== XỬ LÝ VĂN BẢN =====================
def is_question_start(text):
    if not text or len(text) < 5: return False
    if MCQ_OPTION_RE.match(text) or TF_OPTION_RE.match(text): return False
    if ANSWER_LINE_RE.search(text) or ANSWER_BLOCK_RE.match(text): return False
    if LEVEL_RE.search(text): return False
    if re.match(r'^\s*Câu\s*\d+\s*[\.:]\s*', text, re.IGNORECASE): return True
    if re.match(r'^\s*\d+\.\s+', text): return True
    return False

def find_answer_from_line(text):
    match = ANSWER_LINE_RE.search(text)
    return match.group(1).upper() if match else None

def extract_question_number(text):
    match = re.match(r'^\s*(?:Câu\s*(\d+)|(\d+))[\.:]\s*', text, re.IGNORECASE)
    return match.group(1) or match.group(2) if match else None

def clean_question_text(text):
    text = re.sub(r'^(?:Câu\s*\d+|\d+)[\.:]\s*', '', text, flags=re.IGNORECASE)
    return text.strip()

def collect_question_content(paragraphs, start_idx):
    content_lines = []
    i = start_idx
    first_text = paragraphs[i].text
    content_lines.append(clean_question_text(first_text))
    i += 1
    while i < len(paragraphs):
        text = paragraphs[i].text
        stripped = text.strip()
        if not stripped:
            i += 1
            continue
        if MCQ_OPTION_RE.match(stripped) or TF_OPTION_RE.match(stripped):
            break
        if is_question_start(stripped):
            break
        if ANSWER_LINE_RE.search(stripped) or ANSWER_BLOCK_RE.match(stripped):
            break
        content_lines.append(text)
        i += 1
    return "\n".join(content_lines), i

def collect_options(paragraphs, start_idx):
    options = {}
    i = start_idx
    current_label = None
    current_content = []
    current_paras = []  # LƯU TRỰC TIẾP PARAGRAPH
    
    while i < len(paragraphs):
        text = paragraphs[i].text
        stripped = text.strip()
        
        if not stripped:
            i += 1
            continue
            
        if is_question_start(stripped):
            break
        if ANSWER_LINE_RE.search(stripped) or ANSWER_BLOCK_RE.match(stripped):
            break
            
        mcq_match = MCQ_OPTION_RE.match(stripped)
        if mcq_match:
            if current_label and current_content:
                options[current_label] = {'content': '\n'.join(current_content).strip(), 'paras': current_paras}
            current_label = mcq_match.group(1).upper()
            pos = text.find(')') if ')' in text else text.find('.')
            content = text[pos+1:].lstrip() if pos != -1 else ""
            current_content = [content]
            current_paras = [paragraphs[i]]  # CHỐT PARA NGAY TẠI ĐÂY
            i += 1
            continue
            
        tf_match = TF_OPTION_RE.match(stripped)
        if tf_match:
            if current_label and current_content:
                options[current_label] = {'content': '\n'.join(current_content).strip(), 'paras': current_paras}
            current_label = tf_match.group(1).upper()
            pos = text.find(')') if ')' in text else text.find('.')
            content = text[pos+1:].lstrip() if pos != -1 else ""
            current_content = [content]
            current_paras = [paragraphs[i]]  # CHỐT PARA NGAY TẠI ĐÂY
            i += 1
            continue
            
        if current_label:
            current_content.append(text)
            current_paras.append(paragraphs[i]) # LƯU THÊM PARA NẾU ĐÁP ÁN XUỐNG DÒNG
            i += 1
            continue
            
        i += 1
        
    if current_label and current_content:
        options[current_label] = {'content': '\n'.join(current_content).strip(), 'paras': current_paras}
        
    return options, i

def debug_option_format(options):
    """In chi tiết định dạng từng run của tất cả đáp án"""
    for label, opt in options.items():
        para = opt['para']
        print(f"\n----- ĐÁP ÁN {label} -----")
        print(f"Nội dung đầy đủ: {para.text}")
        for idx, run in enumerate(para.runs):
            txt = run.text
            if not txt:
                continue
            # Kiểm tra các thuộc tính
            u = run_is_underlined(run)
            r = run_has_red_color(run)
            m = is_run_marked_as_answer(run)  # nếu còn dùng
            # Lấy màu chữ thực tế
            color_hex = None
            try:
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    if hasattr(rgb, 'rgb'):
                        color_hex = f"{rgb.rgb[0]:02X}{rgb.rgb[1]:02X}{rgb.rgb[2]:02X}"
                    else:
                        color_hex = str(rgb)
            except:
                pass
            
            print(f"  Run {idx}: '{txt}'")
            print(f"    underline = {u} (raw: {run.font.underline})")
            print(f"    red_color = {r} (color: {color_hex})")
            print(f"    xml_mark  = {m}")
            # In thêm XML để kiểm tra thủ công nếu cần (có thể giới hạn độ dài)
            xml_snippet = run._element.xml[:300] if hasattr(run, '_element') else "N/A"
            print(f"    XML snippet: {xml_snippet}")

def find_correct_answer_mcq(options, paragraphs, start_idx):
    # 1. Tìm dòng "Đáp án: X"
    i = start_idx
    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        if is_question_start(text):
            break
        answer = find_answer_from_line(text)
        if answer and answer in options:
            return answer
        i += 1

    # 2. Tìm đáp án dựa trên gạch chân hoặc màu đỏ
    for label, opt in options.items():
        paras = opt.get('paras', [])
        for para in paras:
            for run in para.runs:
                # Hàm is_run_marked_as_answer của bạn viết rất tốt, mình tận dụng luôn
                if is_run_marked_as_answer(run):
                    print(f"   → Chọn đáp án {label} do có định dạng (gạch chân/màu đỏ/XML).")
                    return label
    return None

def find_correct_answer_tf(options, paragraphs, start_idx):
    correct_labels = set()

    # 1. Ưu tiên dòng "Đáp án: A" hoặc tương tự
    i = start_idx
    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        if is_question_start(text): 
            break
        answer = find_answer_from_line(text)
        if answer:
            correct_labels.add(answer.upper())
        i += 1

    # 2. KIỂM TRA TRỰC TIẾP TRÊN PARA ĐÃ THU THẬP
    for label, opt in options.items():
        paras = opt.get('paras', [])
        for para in paras:
            for run in para.runs:
                if is_run_marked_as_answer(run):
                    correct_labels.add(label.upper())
                    print(f"         → TÌM THẤY ĐÁNH DẤU → Đáp án {label}")
                    break  

    if correct_labels:
        result = list(correct_labels)
        print(f"   ✅ TF trả về đáp án: {result}")
        return result
    else:
        return None
def split_inline_options(doc):
    print("\n✂️ Đang tách các đáp án chung dòng (giữ nguyên định dạng)...")
    split_pattern = re.compile(r'\s+([B-D][\.\)])\s*')
    for para in list(doc.paragraphs):
        full_text = para.text
        matches = list(split_pattern.finditer(full_text))
        if not matches: continue
        ranges = []
        last_idx = 0
        for m in matches:
            ranges.append((last_idx, m.start()))
            last_idx = m.start(1)
        ranges.append((last_idx, len(full_text)))
        parent = para._element.getparent()
        current_para_el = para._element
        for start_idx, end_idx in ranges:
            new_p_el = copy.deepcopy(para._element)
            new_para = Paragraph(new_p_el, para._parent)
            current_run_start = 0
            for run in new_para.runs:
                run_len = len(run.text)
                run_end = current_run_start + run_len
                overlap_start = max(current_run_start, start_idx)
                overlap_end = min(run_end, end_idx)
                if overlap_start < overlap_end:
                    local_start = overlap_start - current_run_start
                    local_end = overlap_end - current_run_start
                    run.text = run.text[local_start:local_end]
                else:
                    run.text = ""
                current_run_start = run_end
            current_para_el.addnext(new_p_el)
            current_para_el = new_p_el
        parent.remove(para._element)
    return doc

def split_manual_line_breaks(doc):
    print("✂️ Đang xử lý Shift+Enter (giữ định dạng chính xác)...")
    for para in list(doc.paragraphs):
        full_text = para.text
        if '\n' not in full_text: continue
        lines = full_text.split('\n')
        if len(lines) <= 1: continue
        ranges = []
        last_idx = 0
        for line in lines:
            line_len = len(line)
            ranges.append((last_idx, last_idx + line_len))
            last_idx = last_idx + line_len + 1
        parent = para._element.getparent()
        current_el = para._element
        for start_idx, end_idx in ranges:
            if start_idx == end_idx or not full_text[start_idx:end_idx].strip(): continue
            new_p_el = copy.deepcopy(para._element)
            new_p = Paragraph(new_p_el, para._parent)
            current_run_start = 0
            for run in new_p.runs:
                run_len = len(run.text)
                run_end = current_run_start + run_len
                overlap_start = max(current_run_start, start_idx)
                overlap_end = min(run_end, end_idx)
                if overlap_start < overlap_end:
                    local_start = overlap_start - current_run_start
                    local_end = overlap_end - current_run_start
                    run.text = run.text[local_start:local_end]
                else:
                    run.text = ""
                current_run_start = run_end
            current_el.addnext(new_p_el)
            current_el = new_p_el
        parent.remove(para._element)
    return doc

def process_document(doc):
    """Chỉ tách dòng + lưu DEBUG + reload + xóa level/table.
    Toàn bộ tách câu hỏi & tìm đáp án sẽ chạy trên DEBUG_CLEAN."""
    doc = split_inline_options(doc)
    doc = split_manual_line_breaks(doc)
    
    debug_path = "DEBUG_AFTER_SPLIT.docx"
    doc.save(debug_path)
    print(f"  💾 Đã lưu DEBUG_AFTER_SPLIT.docx (tách dòng OK, giữ nguyên màu/gạch chân)")

    # ==================== RELOAD + XỬ LÝ SẠCH ====================
    doc = Document(debug_path)
    doc = remove_level_lines(doc)
    doc = remove_answer_table(doc)
    
    # LƯU BẢN RELOAD (file thực tế dùng để tách câu hỏi)
    clean_debug_path = "DEBUG_CLEAN.docx"
    doc.save(clean_debug_path)
    print(f"  💾 Đã lưu DEBUG_CLEAN.docx (bản sạch dùng để tách câu hỏi & tìm đáp án)")

    return doc

def extract_questions_to_json(doc, output_json):
    paragraphs = list(doc.paragraphs)
    questions = []
    i = 0
    missing_ids = []

    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        if not text:
            i += 1
            continue
        if is_question_start(text):
            content, next_idx = collect_question_content(paragraphs, i)
            options, opt_end_idx = collect_options(paragraphs, next_idx)

            if not options:
                i = opt_end_idx
                continue

            # Phân loại dựa trên ký tự đầu tiên của nhãn
            first_label = next(iter(options.keys()))
            is_tf = first_label.islower()

            if is_tf:
                correct_list = find_correct_answer_tf(options, paragraphs, opt_end_idx)
                correct = correct_list[0] if correct_list else None
            else:
                correct = find_correct_answer_mcq(options, paragraphs, opt_end_idx)

            q_data = {
                "id": len(questions) + 1,
                "question": content,
                "options": {k: v['content'] for k, v in options.items()},
                "correct": correct
            }
            questions.append(q_data)
            if correct is None:
                missing_ids.append(q_data['id'])
            i = opt_end_idx
        else:
            i += 1

    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)

    if missing_ids:
        print(f"\n⚠️ Cảnh báo: Có {len(missing_ids)} câu không tìm thấy đáp án (ID: {missing_ids})")
    return questions

# ========== HÀM XỬ LÝ FILE WORD (rút gọn) ==========
def process_word_file(input_file,app_dir):
    """Trả về danh sách câu hỏi (dict) và thư mục ảnh tạm."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    image_folder = f"temp_images_{timestamp}"
    
    #temp_file = input_file.replace('.docx', '_temp.docx')
    temp_file = os.path.join(app_dir, 'temp_' + os.path.basename(input_file))
    if not convert_numbering_to_text_with_word(input_file, temp_file):
        raise Exception("Không thể chuyển đổi numbering!")
    
    doc = Document(temp_file)
    doc = extract_images_to_placeholder(doc, image_folder)
    doc = process_document(doc)
    
    try:
        os.remove(temp_file)
    except:
        pass
    
    json_temp = input_file.replace('.docx', '_temp_questions.json')
    extract_questions_to_json(doc, json_temp)
    with open(json_temp, 'r', encoding='utf-8') as f:
        questions = json.load(f)
    os.remove(json_temp)
    
    return questions, image_folder

def convert_to_dataframe(questions):
    """Chuyển list câu hỏi sang DataFrame cho game."""
    rows = []
    for q in questions:
        options = q.get('options', {})
        correct_label = q.get('correct')
        # Chỉ lấy câu hỏi MCQ (đáp án A-D)
        if correct_label not in options or correct_label.islower():
            continue
        labels = ['A', 'B', 'C', 'D']
        if not all(l in options for l in labels):
            continue
        meaning = options[correct_label]
        wrongs = [options[l] for l in labels if l != correct_label]
        if len(wrongs) != 3:
            continue
        rows.append({
            'word': q['question'],
            'meaning': meaning,
            'wrong1': wrongs[0],
            'wrong2': wrongs[1],
            'wrong3': wrongs[2]
        })
    return pd.DataFrame(rows)

# ===================== GIAO DIỆN CUSTOMTKINTER =====================
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

class GameLauncher(ctk.CTk):
    def __init__(self):
        super().__init__()
        # === XÁC ĐỊNH THƯ MỤC GỐC CỦA ỨNG DỤNG (Nuitka onefile) ===
        if sys.argv and os.path.splitext(sys.argv[0])[1].lower() == '.exe':
            # Đang chạy từ file .exe đã build bằng Nuitka
            self.app_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
        else:
            # Đang chạy từ mã nguồn Python
            self.app_dir = os.path.dirname(os.path.abspath(__file__))
        self.title("🚀 Ứng dụng Học Từ Vựng Tiếng Anh")
        self.geometry("800x450")
        self.resizable(False, False)
        
        self.word_file = ctk.StringVar()
        self.status_text = ctk.StringVar(value="Chưa chọn file Word")
        self.server_process = None
        self.image_folder = None
        
        self.create_widgets()
        self.center_window()
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
    
    def center_window(self):
        self.update_idletasks()
        w, h = 800, 450
        x = (self.winfo_screenwidth() // 2) - (w // 2)
        y = (self.winfo_screenheight() // 2) - (h // 2)
        self.geometry(f'{w}x{h}+{x}+{y}')
    
    def create_widgets(self):
        # Tiêu đề
        title_label = ctk.CTkLabel(self, text="📚 HỌC TỪ VỰNG TIẾNG ANH QUA TRÒ CHƠI", 
                                   font=ctk.CTkFont(size=20, weight="bold"))
        title_label.pack(pady=20)
        
        # Khung chọn file
        file_frame = ctk.CTkFrame(self)
        file_frame.pack(pady=10, padx=20, fill="x")
        
        ctk.CTkLabel(file_frame, text="📄 File Word câu hỏi:", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, padx=10, pady=10, sticky="w")
        entry = ctk.CTkEntry(file_frame, textvariable=self.word_file, width=350, placeholder_text="Chọn file .docx")
        entry.grid(row=0, column=1, padx=5, pady=10)
        ctk.CTkButton(file_frame, text="Duyệt...", width=80, command=self.browse_file).grid(row=0, column=2, padx=10, pady=10)
        
        # Nút xử lý và bắt đầu
        self.process_btn = ctk.CTkButton(self, text="🎮 XỬ LÝ & BẮT ĐẦU HỌC", 
                                         command=self.process_and_launch, 
                                         width=200, height=40,
                                         font=ctk.CTkFont(size=14, weight="bold"))
        self.process_btn.pack(pady=20)
        
        # Thanh trạng thái
        status_frame = ctk.CTkFrame(self)
        status_frame.pack(pady=10, padx=20, fill="x")
        ctk.CTkLabel(status_frame, textvariable=self.status_text, anchor="w").pack(fill="x", padx=10, pady=5)
        
        # >>> THÊM LABEL HIỂN THỊ IP Ở ĐÂY <<<
        self.ip_label = ctk.CTkLabel(
            self, 
            text="", 
            font=ctk.CTkFont(size=13, weight="bold"),
            text_color="#2ecc71"
        )
        self.ip_label.pack(pady=5)
        
        # Hướng dẫn
        info_text = ("1. Chọn file Word (.docx) chứa câu hỏi trắc nghiệm (MCQ).\n"
                     "2. Nhấn nút trên để trích xuất và lưu vào questions.xlsx.\n"
                     "3. Game sẽ tự động khởi động trong trình duyệt.")
        info_label = ctk.CTkLabel(self, text=info_text, justify="left", font=ctk.CTkFont(size=12))
        info_label.pack(pady=10)
    
    def browse_file(self):
        filename = filedialog.askopenfilename(
            title="Chọn file Word",
            filetypes=[("Word files", "*.docx")]
        )
        if filename:
            self.word_file.set(filename)
            self.status_text.set(f"Đã chọn: {os.path.basename(filename)}")
    
    def process_and_launch(self):
        if not self.word_file.get():
            messagebox.showerror("Lỗi", "Vui lòng chọn file Word trước.")
            return
        
        self.process_btn.configure(state="disabled")
        self.status_text.set("⏳ Đang xử lý file Word...")
        self.update()
        
        try:
            # 1. Xử lý file Word
            questions, img_folder = process_word_file(self.word_file.get(),self.app_dir)
            self.image_folder = img_folder
            self.status_text.set(f"✅ Đã trích xuất {len(questions)} câu hỏi. Đang chuyển đổi...")
            
            # 2. Chuyển thành DataFrame và lưu Excel
            df = convert_to_dataframe(questions)
            if df.empty:
                raise ValueError("Không có câu hỏi MCQ hợp lệ (cần đủ 4 đáp án và xác định được đáp án đúng).")
            
            df.to_excel(QUESTIONS_FILE, index=False)
            self.status_text.set(f"💾 Đã lưu {len(df)} câu hỏi vào questions.xlsx")
            
            # 3. Dọn file tạm
            if self.image_folder and os.path.exists(self.image_folder):
                shutil.rmtree(self.image_folder)
                self.image_folder = None
            # Xóa các file debug tạm
            debug_files = ["DEBUG_AFTER_SPLIT.docx", "DEBUG_CLEAN.docx"]
            for f in debug_files:
                if os.path.exists(f):
                    try:
                        os.remove(f)
                        self.status_text.set(f"🧹 Đã xóa file tạm: {f}")
                    except Exception as e:
                        self.status_text.set(f"⚠️ Không thể xóa {f}: {e}")
            
            # 4. Kiểm tra và giải phóng cổng 5000
            self.free_port_5000()
            
            # 5. Khởi động Flask server
            self.status_text.set("🚀 Đang khởi động máy chủ game...")
            self.start_flask_server()
            
            # 6. Chờ server sẵn sàng và mở trình duyệt
            # Hiển thị IP cho máy khác
            local_ip = get_local_ip()
            self.ip_label.configure(text=f"🌐 Máy khác truy cập: http://{local_ip}:5000")
            self.status_text.set(f"Server đã sẵn sàng!")
            #self.status_text.set(f"🌐 Máy khác có thể truy cập tại: http://{local_ip}:5000")
            self.after(5000, lambda: webbrowser.open(f'http://{local_ip}:5000'))
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Xử lý thất bại:\n{str(e)}")
            self.status_text.set(f"❌ Lỗi: {e}")
            self.process_btn.configure(state="normal")
    
    def free_port_5000(self):
        """Giải phóng cổng 5000 triệt để."""
        try:
            if os.name == 'nt':
                # Windows: tìm PID nghe cổng 5000 và kill
                result = subprocess.run(
                    'for /f "tokens=5" %a in (\'netstat -ano ^| findstr :5000 ^| findstr LISTENING\') do taskkill /F /PID %a',
                    shell=True, capture_output=True, text=True
                )
                if result.returncode == 0:
                    print("[LAUNCHER] Đã giải phóng cổng 5000")
                else:
                    print("[LAUNCHER] Cổng 5000 đã trống hoặc không cần giải phóng")
            else:
                subprocess.run("lsof -ti:5000 | xargs kill -9", shell=True)
        except Exception as e:
            print(f"[LAUNCHER] Lỗi giải phóng cổng: {e}")
    
    def monitor_server_output(self):
        """Đọc output từ server và phát hiện khi server dừng."""
        for line in self.server_process.stdout:
            print(f"[SERVER] {line.strip()}")
        
        # Khi vòng lặp kết thúc (server đã tắt)
        print("[LAUNCHER] Server process đã kết thúc")
        self.server_process = None
        self.after(0, lambda: self.status_text.set("⏹️ Server đã dừng"))
    
    def start_flask_server(self):
        """Khởi động app.py hoặc app.exe như một tiến trình con."""
        # Ưu tiên tìm file .exe trước (khi đã build)
        app_exe = os.path.join(self.app_dir, "app.exe")
        app_py = os.path.join(self.app_dir, "app.py")
        
        if os.path.exists(app_exe):
            cmd = [app_exe]
            self.status_text.set("🔧 Sử dụng file app.exe đã build.")
        elif os.path.exists(app_py):
            cmd = [sys.executable, app_py]
            self.status_text.set("🐍 Chạy app.py bằng Python interpreter.")
        else:
            raise FileNotFoundError(
                f"Không tìm thấy app.exe hoặc app.py trong:\n"
                f"- {app_exe}\n- {app_py}"
            )
        
        creationflags = subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0
        self.server_process = subprocess.Popen(
            cmd,
            cwd=self.app_dir,  # Quan trọng: đặt thư mục làm việc
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            encoding='utf-8',
            errors='replace',
            creationflags=creationflags
        )
        threading.Thread(target=self.monitor_server_output, daemon=True).start()
    
    def monitor_server_output(self):
        """Đọc output từ server và ghi vào console (không hiển thị lên GUI để tránh lag)."""
        for line in self.server_process.stdout:
            print(f"[SERVER] {line.strip()}")
    
    def on_closing(self):
        """Dừng server triệt để khi đóng cửa sổ."""
        if self.server_process:
            self.status_text.set("🛑 Đang dừng server...")
            
            # 1. Thử terminate nhẹ nhàng trước
            self.server_process.terminate()
            try:
                self.server_process.wait(timeout=2)
                print("[LAUNCHER] Server đã dừng bằng terminate()")
            except subprocess.TimeoutExpired:
                # 2. Nếu không chịu dừng, kill mạnh tay
                print("[LAUNCHER] Server không dừng, đang kill()...")
                self.server_process.kill()
                try:
                    self.server_process.wait(timeout=2)
                except subprocess.TimeoutExpired:
                    pass
            
            # 3. Kiểm tra lại xem còn sống không
            if self.server_process.poll() is None:
                print("[LAUNCHER] Server vẫn còn sống, dùng taskkill...")
                self.force_kill_app_exe()
            
            self.server_process = None
        
        # 4. Giải phóng cổng 5000 lần cuối
        self.free_port_5000()
        
        self.destroy()

    def force_kill_app_exe(self):
        """Dùng lệnh hệ thống để kill triệt để app.exe."""
        try:
            if os.name == 'nt':
                # Windows: kill tất cả app.exe (chỉ trong thư mục hiện tại thì an toàn hơn)
                subprocess.run('taskkill /F /IM app.exe', shell=True, capture_output=True)
                print("[LAUNCHER] Đã taskkill /F /IM app.exe")
            else:
                # Linux/macOS
                subprocess.run("pkill -f app.exe", shell=True)
        except Exception as e:
            print(f"[LAUNCHER] Lỗi force kill: {e}")

if __name__ == "__main__":
    app = GameLauncher()
    app.mainloop()